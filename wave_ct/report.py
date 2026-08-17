"""Create a dataset-driven CT inversion report without assumed anomaly geometry."""

from __future__ import annotations

import argparse
import csv
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def read_stats(path: Path) -> dict[str, str]:
    values: dict[str, str] = {}
    if not path.is_file():
        return values
    for line in path.read_text(encoding="utf-8").splitlines():
        if "=" in line and not line.startswith("="):
            key, value = line.split("=", 1)
            values[key.strip()] = value.strip()
    return values


def csv_summary(path: Path) -> tuple[int, int, int]:
    rows = list(csv.DictReader(path.open("r", encoding="utf-8-sig", newline="")))
    sources = {row.get("震源编号", "") for row in rows}
    stations = {
        (row.get("台站坐标-x", ""), row.get("台站坐标-y", ""), row.get("台站坐标-z", ""))
        for row in rows
    }
    return len(rows), len(sources), len(stations)


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def set_cell(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)


def add_picture(document: Document, path: Path, caption: str, width: float = 16.8) -> None:
    if not path.is_file():
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(caption)
    run.bold = True
    run.font.size = Pt(10.5)
    document.add_picture(str(path), width=Cm(width))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_pair(document: Document, left: Path, right: Path, left_title: str, right_title: str) -> None:
    if not left.is_file() and not right.is_file():
        return
    table = document.add_table(rows=2, cols=2)
    for cell, title in zip(table.rows[0].cells, (left_title, right_title)):
        set_cell(cell, title, bold=True)
    for cell, path in zip(table.rows[1].cells, (left, right)):
        if path.is_file():
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.add_run().add_picture(str(path), width=Cm(8.0))


def first_existing(root: Path, *names: str) -> Path:
    """Find an asset in the stage root or in a nested final-results folder."""
    for name in names:
        direct = root / name
        if direct.is_file():
            return direct
        matches = sorted(root.rglob(name))
        if matches:
            return matches[0]
    return root / names[0]


def all_assets(root: Path, pattern: str) -> list[Path]:
    seen: set[str] = set()
    result: list[Path] = []
    for path in sorted(root.rglob(pattern)):
        # Rendering may mirror the same deliverable into a compatibility
        # subfolder.  A report needs one copy per semantic filename.
        key = path.name.lower()
        if key not in seen:
            seen.add(key)
            result.append(path)
    return result


def main() -> None:
    parser = argparse.ArgumentParser(description="生成通用震动波CT反演报表")
    parser.add_argument("--dataset-name", required=True)
    parser.add_argument("--input-csv", type=Path, required=True)
    parser.add_argument("--inversion-dir", type=Path, required=True)
    parser.add_argument("--validation-dir", type=Path, required=True)
    parser.add_argument("--output-docx", type=Path, required=True)
    parser.add_argument("--picking-source", required=True)
    parser.add_argument("--note", action="append", default=[])
    args = parser.parse_args()

    stats = read_stats(args.inversion_dir / "slice_report.txt")
    rows, sources, stations = csv_summary(args.input_csv)
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)
    for style_name in ("Normal", "Title", "Heading 1", "Heading 2"):
        style = document.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    document.add_paragraph("\n\n")
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(args.dataset_name)
    run.bold = True
    run.font.size = Pt(22)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Wave CT Studio 反演与软件验收报表")
    run.bold = True
    run.font.size = Pt(26)
    document.add_page_break()

    document.add_heading("一、输入与计算概况", level=1)
    table = document.add_table(rows=2, cols=8)
    table.style = "Table Grid"
    headers = ["输入射线", "震源", "台站", "反演射线", "训练事件", "验证事件", "验证RMS", "可靠覆盖"]
    values = [
        str(rows), str(sources), str(stations), stats.get("n_rays", "-"),
        stats.get("training_sources", "-"), stats.get("validation_sources", "-"),
        f"{float(stats.get('best_validation_rms_s', 0.0)) * 1000.0:.2f} ms",
        f"{float(stats.get('reliable_coverage_fraction', 0.0)) * 100.0:.2f}%",
    ]
    for index, value in enumerate(headers):
        set_cell(table.rows[0].cells[index], value, bold=True)
        set_cell_shading(table.rows[0].cells[index], "D9EAF7")
    for index, value in enumerate(values):
        set_cell(table.rows[1].cells[index], value)
    document.add_paragraph(f"P波到时来源：{args.picking_source}")
    document.add_paragraph(f"输入文件：{args.input_csv}")

    document.add_heading("二、P波到时质量检查", level=1)
    add_pair(
        document,
        args.validation_dir / "01_表观速度直方图.png",
        args.validation_dir / "02_距离到时散点图.png",
        "表观速度分布", "距离-到时关系",
    )
    add_picture(document, args.validation_dir / "03_典型波形拾取检查.png", "典型P波拾取位置抽查")
    add_picture(
        document,
        args.validation_dir / "01_W2人工P标记质控.png",
        "W2文件头人工P标记的事件中心化速度质控",
    )
    add_picture(
        document,
        args.validation_dir / "02_W2人工P标记波形抽查.png",
        "直接解码W2原始通道后的人工P标记抽查",
    )

    document.add_heading("三、反演收敛与射线覆盖", level=1)
    add_picture(document, args.inversion_dir / "rms_convergence.png", "训练集与按震源留出验证集RMS")
    for diagnostic in (
        first_existing(args.validation_dir, "01_走时残差诊断.png"),
        first_existing(args.validation_dir, "02_异常体XZ_YZ连续性.png"),
    ):
        add_picture(document, diagnostic, diagnostic.stem)

    coverage_images = all_assets(args.inversion_dir, "ray_coverage_z*.png")
    for path in coverage_images:
        add_picture(document, path, f"射线覆盖：{path.stem.replace('ray_coverage_', '')}")

    document.add_heading("四、P波速度切片", level=1)
    velocity_images = all_assets(args.inversion_dir, "velocity_z*.png")
    for path in velocity_images:
        add_picture(document, path, f"P波速度切片：{path.stem.replace('velocity_', '')}")

    source_map = first_existing(args.inversion_dir, "source_distribution.png")
    ray_plan = first_existing(args.inversion_dir, "ray_coverage_plan.png")
    ray_3d = first_existing(args.inversion_dir, "ray_coverage_3d.png")
    workface_velocity = all_assets(args.inversion_dir, "key_workface_velocity_z*.png")
    if not workface_velocity:
        workface_velocity = all_assets(args.inversion_dir, "surfer_style_velocity_z*.png")
    anomaly_maps = all_assets(args.inversion_dir, "key_workface_anomaly_z*.png")
    if source_map.is_file() or ray_plan.is_file() or workface_velocity or anomaly_maps:
        document.add_heading("五、工作面专题成果", level=1)
        add_pair(document, source_map, ray_plan, "震源分布图", "反演射线覆盖图")
        add_picture(document, ray_3d, "三维反演射线覆盖图")
        for path in workface_velocity:
            elevation = path.stem.replace("yanbei_velocity_z", "").replace(
                "surfer_style_velocity_z", ""
            )
            label = "平滑展示图" if path.stem.startswith("surfer_style") else "CT反演结果图"
            add_picture(document, path, f"{elevation}标高工作面煤层{label}")
        for path in anomaly_maps:
            elevation = path.stem.replace("anomaly_z", "")
            add_picture(document, path, f"{elevation}标高波速异常系数An分布图")

    comparison_dir = args.inversion_dir / "参考结果对比"
    for path in all_assets(
        args.inversion_dir, "key_workface_anomaly_legacy_disabled_z*.png"
    ):
        add_picture(
            document,
            path,
            "弱异常增强显示（仅观察空间形态；异常强度以固定色标定量图为准）",
        )

    comparison_images = sorted(comparison_dir.glob("same_basemap_comparison_z*.png"))
    comparison_images.extend(sorted(comparison_dir.glob("comparison_z*.png")))
    if comparison_images:
        document.add_heading("六、与外部参考结果的独立对比", level=1)
        document.add_paragraph(
            "下列参考TXT数值未参与反演求解，只在反演完成后进行同坐标逐点评价。"
            "图中的差异不能通过修改观测或写入参考异常体来消除。"
        )
        for path in comparison_images:
            add_picture(document, path, path.stem.replace("comparison_z", "标高 ") + " m 数值对比")
        summary_path = comparison_dir / "comparison_summary.txt"
        if summary_path.is_file():
            paragraph = document.add_paragraph()
            run = paragraph.add_run(summary_path.read_text(encoding="utf-8", errors="replace"))
            run.font.name = "Consolas"
            run.font.size = Pt(8.5)

    document.add_heading("七、质量结论与解释边界", level=1)
    cad_record = first_existing(args.inversion_dir, "cad_background_source.txt")
    cad_source = "none"
    if cad_record.is_file():
        for line in cad_record.read_text(encoding="utf-8", errors="replace").splitlines():
            if line.startswith("source="):
                cad_source = line.split("=", 1)[1].strip()
                break
    bullets = [
        "反演只使用输入CSV中的观测走时、震源坐标和台站坐标，未写入预期异常体。",
        f"验证方式：{stats.get('validation_split', '按反演日志')}；验证指标：{stats.get('validation_metric', '-')}。",
        f"可正演求解模型速度范围：{stats.get('velocity_range', '-')} m/s；低覆盖显示场范围：{stats.get('display_velocity_range', '-')} m/s；背景速度：{stats.get('background_velocity', '-')} m/s。",
        f"三维路径覆盖率：{float(stats.get('covered_fraction', 0.0)) * 100.0:.2f}%；可靠覆盖率：{float(stats.get('reliable_coverage_fraction', 0.0)) * 100.0:.2f}%。可靠覆盖边界外不能解释为已验证异常。",
        f"导出求解模型验证RMS：{float(stats.get('solution_validation_rms_s', stats.get('best_validation_rms_s', 0.0))) * 1000.0:.2f} ms；显示场仅用于低覆盖区稳定展示，不用于正演精度声明。",
        "速度异常应与射线覆盖、拾取质量和独立地质资料共同解释。",
    ] + args.note
    if cad_source and cad_source != "none":
        bullets.append(f"矿井底图来源：{cad_source}；底图只用于坐标叠加展示，不参与CT求解。")
    elif source_map.is_file() or workface_velocity:
        bullets.append("本次工作面专题图未叠加矿井底图，只显示边界、震源、台站和反演场。")
    for text in bullets:
        document.add_paragraph(text, style="List Bullet")

    args.output_docx.parent.mkdir(parents=True, exist_ok=True)
    document.save(args.output_docx)
    print(f"saved {args.output_docx}")


if __name__ == "__main__":
    main()
