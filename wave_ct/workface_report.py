"""Create a traceable Yanbei workface CT report in the reference layout."""

from __future__ import annotations

import argparse
import json
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def read_key_values(path: Path) -> dict[str, str]:
    values: dict[str, str] = {}
    for line in path.read_text(encoding="utf-8").splitlines():
        if "=" in line and not line.startswith("="):
            key, value = line.split("=", 1)
            values[key.strip()] = value.strip()
    return values


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def set_cell_text(cell, text: str, bold: bool = False, size: float = 9.0) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)


def add_stats_table(document: Document, summary: dict[str, object], inversion: dict[str, str]) -> None:
    table = document.add_table(rows=2, cols=8)
    table.style = "Table Grid"
    headers = ["震动总数", "反演震动数", "反演台站数", "射线数", "反演后射线", "初始RMS", "验证RMS", "可靠覆盖"]
    values = [
        str(summary.get("catalog_events_in_period", "-")),
        str(summary.get("inversion_events", "-")),
        str(summary.get("unique_station_count", "-")),
        str(summary.get("inversion_rays", "-")),
        inversion.get("n_rays", "-"),
        f"{float(inversion.get('initial_rms_s', 0.0)) * 1000.0:.2f} ms",
        f"{float(inversion.get('best_validation_rms_s', 0.0)) * 1000.0:.2f} ms",
        f"{float(inversion.get('reliable_coverage_fraction', 0.0)) * 100.0:.2f}%",
    ]
    for index, text in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], text, bold=True)
        set_cell_shading(table.rows[0].cells[index], "D9EAF7")
    for index, text in enumerate(values):
        set_cell_text(table.rows[1].cells[index], text, size=9.5)


def add_picture(document: Document, path: Path, caption: str, width_cm: float = 16.8) -> None:
    if not path.exists():
        paragraph = document.add_paragraph(f"缺失图件：{path.name}")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(caption)
    run.bold = True
    run.font.size = Pt(11)
    document.add_picture(str(path), width=Cm(width_cm))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_two_pictures(document: Document, left: Path, right: Path, left_caption: str, right_caption: str) -> None:
    table = document.add_table(rows=2, cols=2)
    table.autofit = False
    for cell, caption in zip(table.rows[0].cells, (left_caption, right_caption)):
        set_cell_text(cell, caption, bold=True, size=9.5)
    for cell, path in zip(table.rows[1].cells, (left, right)):
        if path.exists():
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.add_run().add_picture(str(path), width=Cm(8.0))
        else:
            set_cell_text(cell, f"缺失：{path.name}")


def set_document_fonts(document: Document) -> None:
    for style_name in ("Normal", "Title", "Heading 1", "Heading 2"):
        style = document.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    document.styles["Normal"].font.size = Pt(10.5)


def add_cover(document: Document, period: str) -> None:
    document.add_paragraph("\n\n\n")
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("砚北煤矿150404工作面回采期间")
    run.bold = True
    run.font.size = Pt(24)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Wave CT Studio 震动波 CT 反演报表")
    run.bold = True
    run.font.size = Pt(30)
    period_paragraph = document.add_paragraph()
    period_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = period_paragraph.add_run(period)
    run.font.size = Pt(16)
    document.add_paragraph("\n\n\n\n")
    organization = document.add_paragraph()
    organization.alignment = WD_ALIGN_PARAGRAPH.CENTER
    organization.add_run("Wave CT Studio").font.size = Pt(16)
    document.add_page_break()


def main() -> None:
    parser = argparse.ArgumentParser(description="生成砚北四期CT DOCX报表")
    parser.add_argument("summary_json", type=Path)
    parser.add_argument("slice_report", type=Path)
    parser.add_argument("image_dir", type=Path)
    parser.add_argument("output_docx", type=Path)
    parser.add_argument("--period", required=True)
    parser.add_argument("--slice-z", required=True)
    args = parser.parse_args()

    summary = json.loads(args.summary_json.read_text(encoding="utf-8"))
    inversion = read_key_values(args.slice_report)
    slices = [int(float(value.strip())) for value in args.slice_z.split(",") if value.strip()]
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)
    set_document_fonts(document)
    add_cover(document, args.period)

    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run(f"砚北煤矿150404工作面回采期间震动波CT反演报表（{args.period}）")
    run.bold = True
    run.font.size = Pt(15)
    add_stats_table(document, summary, inversion)

    requested_start = summary.get("requested_start")
    available_start = summary.get("available_picked_start")
    requested_end = summary.get("requested_end")
    available_end = summary.get("available_picked_end")
    missing_start = (
        requested_start and available_start
        and datetime.fromisoformat(str(available_start)).date()
        > datetime.fromisoformat(str(requested_start)).date()
    )
    missing_end = (
        requested_end and available_end
        and datetime.fromisoformat(str(available_end)).date()
        < datetime.fromisoformat(str(requested_end)).date()
    )
    if missing_start or missing_end:
        paragraph = document.add_paragraph()
        run = paragraph.add_run("数据完整性提示：")
        run.bold = True
        details = []
        if missing_start:
            details.append(f"申请周期开始于 {requested_start}，首个可用人工P标记事件为 {available_start}")
        if missing_end:
            details.append(f"申请周期截止到 {requested_end}，末个可用人工P标记事件为 {available_end}")
        paragraph.add_run("；".join(details) + "。本报表只反映实际可用数据范围，未补造缺失事件。")

    document.add_heading("一、震源分布与射线覆盖", level=1)
    add_two_pictures(
        document,
        args.image_dir / "source_distribution.png",
        args.image_dir / "ray_coverage_plan.png",
        f"{args.period}期间震源分布图",
        "反演射线覆盖图",
    )
    add_picture(
        document,
        args.image_dir / "ray_coverage_3d.png",
        "三维反演射线覆盖图",
    )

    document.add_heading("二、水平切面P波速度", level=1)
    for index, target in enumerate(slices, start=1):
        add_picture(
            document,
            args.image_dir / f"yanbei_velocity_z{target}.png",
            f"图{index}  {target}标高切面波速分布云图",
        )

    document.add_heading("三、波速异常系数与变化梯度", level=1)
    for target in slices:
        add_two_pictures(
            document,
            args.image_dir / f"anomaly_z{target}.png",
            args.image_dir / f"gradient_z{target}.png",
            f"{target}标高波速异常系数 An",
            f"{target}标高波速变化梯度 VG",
        )

    document.add_page_break()
    document.add_heading("四、垂直切面P波速度", level=1)
    vertical = [
        ("vertical_material_roadway.png", "材料顺槽垂直切面波速分布云图"),
        ("vertical_transport_roadway.png", "运输顺槽垂直切面波速分布云图"),
        ("vertical_workface_center.png", "工作面中线垂直切面波速分布云图"),
        ("vertical_mining_dip.png", "沿150404工作面回采位置倾向切面波速分布云图"),
    ]
    for offset, (filename, caption) in enumerate(vertical, start=len(slices) + 1):
        add_picture(document, args.image_dir / filename, f"图{offset}  {caption}")

    cad_source = read_key_values(args.image_dir / "cad_background_source.txt").get("source", "none")
    document.add_heading("五、质量控制与解释边界", level=1)
    bullets = [
        "P波到时来自原始.W文件头中的人工标记样点，未用距离/3600m/s生成观测走时。",
        f"可正演求解模型范围：{inversion.get('velocity_range', '-')} m/s；低覆盖显示场范围：{inversion.get('display_velocity_range', '-')} m/s；背景速度：{inversion.get('background_velocity', '-')} m/s。",
        f"初始RMS为{float(inversion.get('initial_rms_s', 0))*1000:.2f}ms，最佳验证RMS为{float(inversion.get('best_validation_rms_s', 0))*1000:.2f}ms。",
        f"验证采用按完整震源事件留出：训练事件{inversion.get('training_sources', '-')}个，"
        f"验证事件{inversion.get('validation_sources', '-')}个；验证事件未参与速度模型和震源时刻校正求解。",
        "验证RMS按事件消除一个公共到时偏移后计算，用于检验相对走时所约束的空间速度结构，不能解释为绝对发震时刻误差。",
        f"三维路径覆盖率为{float(inversion.get('covered_fraction', 0))*100:.2f}%，可靠覆盖率为{float(inversion.get('reliable_coverage_fraction', 0))*100:.2f}%；可靠覆盖边界外不能解释为已验证异常。",
        "成果图以虚线标出可靠覆盖边界；边界外的彩色场只用于保持底图连续和显示总体趋势，不作为煤层异常结论。",
        f"导出求解模型验证RMS为{float(inversion.get('solution_validation_rms_s', inversion.get('best_validation_rms_s', 0)))*1000:.2f}ms；覆盖稳定显示场不参与正演精度声明。",
        "矿图线条来自对应月份原始SOS/Mapa.dat；V1-V6工作面多边形来自前人坐标文件，尚未与真实DWG图层及坐标系逐点核验，只能作为近似边界。",
        "速度异常必须结合采掘进度、巷道、断层、钻探和相邻月份变化共同解释，不能仅凭颜色作地质结论。",
    ]
    if cad_source.startswith(("DWG:", "DXF:")):
        bullets[-2] = f"\u77ff\u56fe\u7ebf\u6761\u6765\u6e90\uff1a{cad_source}\uff1bV1-V6\u8fb9\u754c\u5df2\u4e0e\u5e26\u5750\u6807\u5de5\u4f5c\u9762\u56fe\u6838\u5bf9\uff0c\u4ecd\u5e94\u4ee5\u77ff\u65b9\u6700\u65b0CAD\u7248\u672c\u4e3a\u51c6\u3002"

    for text in bullets:
        document.add_paragraph(text, style="List Bullet")

    args.output_docx.parent.mkdir(parents=True, exist_ok=True)
    document.save(args.output_docx)
    print(f"saved {args.output_docx}")


if __name__ == "__main__":
    main()
